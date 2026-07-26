from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Mm, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "docs" / "diplom_revised.docx"
SCREENS = ROOT / "runtime-logs" / "diploma-screens"

ACCENT = "4F6A5D"
BORDER = "C8D2CA"
SHADE = "EEF2EE"
SOFT = "F7F4EF"


def font(run, size=14, name="Times New Roman", bold=False, italic=False, color=None):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def add_field(paragraph, instr: str) -> None:
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = instr
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instr_text, separate, end])


def set_update_fields(doc: Document) -> None:
    settings = doc.settings.element
    update_fields = settings.find(qn("w:updateFields"))
    if update_fields is None:
        update_fields = OxmlElement("w:updateFields")
        settings.append(update_fields)
    update_fields.set(qn("w:val"), "true")


def configure(doc: Document) -> None:
    section = doc.sections[0]
    section.page_width = Mm(210)
    section.page_height = Mm(297)
    section.left_margin = Mm(30)
    section.right_margin = Mm(15)
    section.top_margin = Mm(20)
    section.bottom_margin = Mm(20)
    section.different_first_page_header_footer = True

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Times New Roman"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    normal.font.size = Pt(14)
    normal.paragraph_format.first_line_indent = Cm(1.25)
    normal.paragraph_format.line_spacing = 1.5
    normal.paragraph_format.space_after = Pt(0)
    normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    for style_name in ("Heading 1", "Heading 2", "Heading 3"):
        style = styles[style_name]
        style.font.name = "Times New Roman"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
        style.font.size = Pt(14)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(ACCENT)
        style.paragraph_format.first_line_indent = Cm(0)
        style.paragraph_format.line_spacing = 1.5
        style.paragraph_format.keep_with_next = True
        style.paragraph_format.space_before = Pt(10 if style_name == "Heading 1" else 6)
        style.paragraph_format.space_after = Pt(6)
    styles["Heading 1"].paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_field(footer, "PAGE")
    set_update_fields(doc)


def para(doc: Document, text: str = "", *, align=None, indent=True, bold=False) -> None:
    p = doc.add_paragraph()
    p.alignment = align or WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.first_line_indent = Cm(1.25 if indent else 0)
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run(text)
    font(r, bold=bold)


def center(doc: Document, text: str = "", *, size=14, bold=False, after=0) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = Cm(0)
    p.paragraph_format.line_spacing = 1.15
    p.paragraph_format.space_after = Pt(after)
    r = p.add_run(text)
    font(r, size=size, bold=bold)


def h1(doc: Document, title: str, *, page_break=True) -> None:
    if page_break:
        doc.add_page_break()
    p = doc.add_paragraph(style="Heading 1")
    p.paragraph_format.first_line_indent = Cm(0)
    p.add_run(title.upper())


def h2(doc: Document, title: str) -> None:
    p = doc.add_paragraph(style="Heading 2")
    p.add_run(title)


def caption(doc: Document, text: str, *, left=False) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT if left else WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = Cm(0)
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run(text)
    font(r, size=12, italic=True)


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, value=90) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin in ("top", "start", "bottom", "end"):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_borders(table, color=BORDER) -> None:
    borders = table._tbl.tblPr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        table._tbl.tblPr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        element = borders.find(qn(f"w:{edge}"))
        if element is None:
            element = OxmlElement(f"w:{edge}")
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "4")
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def table(doc: Document, title: str, headers: list[str], rows: list[list[str]]) -> None:
    caption(doc, title, left=True)
    t = doc.add_table(rows=1, cols=len(headers))
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_borders(t)
    for i, header in enumerate(headers):
        cell = t.rows[0].cells[i]
        set_cell_shading(cell, SHADE)
        set_cell_margins(cell)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.first_line_indent = Cm(0)
        r = p.add_run(header)
        font(r, size=11, bold=True)
    for index, row in enumerate(rows):
        cells = t.add_row().cells
        for i, value in enumerate(row):
            cell = cells[i]
            if index % 2:
                set_cell_shading(cell, SOFT)
            set_cell_margins(cell)
            p = cell.paragraphs[0]
            p.paragraph_format.first_line_indent = Cm(0)
            p.paragraph_format.line_spacing = 1.15
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            r = p.add_run(value)
            font(r, size=10.5)
    doc.add_paragraph()


def image(doc: Document, path: Path, cap: str, width=15.0) -> None:
    if not path.exists():
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = Cm(0)
    p.add_run().add_picture(str(path), width=Cm(width))
    caption(doc, cap)


def read(path: str) -> str:
    return (ROOT / path).read_text(encoding="utf-8")


def code_excerpt(doc: Document, title: str, text: str, max_lines=36) -> None:
    caption(doc, title, left=True)
    lines = text.strip("\n").splitlines()[:max_lines]
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Cm(0)
    p.paragraph_format.line_spacing = 1.0
    p.paragraph_format.space_after = Pt(6)
    for index, line in enumerate(lines):
        if index:
            p.add_run("\n")
        r = p.add_run(line[:112])
        font(r, size=10, name="Courier New")


def excerpt_between(text: str, start: str, end: str | None = None) -> str:
    s = text.find(start)
    if s < 0:
        return text[:1400]
    if end is None:
        return text[s : s + 1600]
    e = text.find(end, s + len(start))
    return text[s:e] if e > s else text[s : s + 1600]


def cover(doc: Document) -> None:
    center(doc, "Министерство науки и высшего образования РФ", size=12)
    center(doc, "Федеральное государственное автономное образовательное учреждение высшего образования", size=12)
    center(doc, "«Казанский (Приволжский) Федеральный Университет»", size=12)
    center(doc, "")
    center(doc, "ИНСТИТУТ ВЫЧИСЛИТЕЛЬНОЙ МАТЕМАТИКИ И", size=12, bold=True)
    center(doc, "ИНФОРМАЦИОННЫХ ТЕХНОЛОГИЙ", size=12, bold=True)
    center(doc, "")
    center(doc, "КАФЕДРА ТЕОРЕТИЧЕСКОЙ КИБЕРНЕТИКИ", size=12)
    center(doc, "Направление: 01.03.02. Прикладная математика и информатика", size=12)
    center(doc, "Профиль: Прикладная математика и информатика", size=12)
    for _ in range(4):
        center(doc, "")
    center(doc, "ВЫПУСКНАЯ КВАЛИФИКАЦИОННАЯ РАБОТА", size=14, bold=True, after=8)
    center(
        doc,
        "Разработка веб-системы подбора жаккардовых обоев с модулем нейросетевой визуализации интерьера",
        size=14,
        bold=True,
        after=10,
    )
    for _ in range(3):
        center(doc, "")
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.paragraph_format.first_line_indent = Cm(0)
    r = p.add_run(
        "Работа завершена:\n"
        "Студент гр. _________\n"
        "«__»__________2026 г.                 ____________________\n\n"
        "Работа допущена к защите:\n"
        "Научный руководитель\n"
        "«__»__________2026 г.                 ____________________\n\n"
        "Заведующий кафедрой\n"
        "«__»__________2026 г.                 ____________________"
    )
    font(r)
    for _ in range(5):
        center(doc, "")
    center(doc, "Казань – 2026", size=14)
    doc.add_page_break()


def toc(doc: Document) -> None:
    h1(doc, "СОДЕРЖАНИЕ", page_break=False)
    items = [
        "ВВЕДЕНИЕ",
        "1 Анализ предметной области и требований",
        "1.1 Характеристика предметной области",
        "1.2 Требования к системе",
        "1.3 Выбор технологического стека",
        "2 Проектирование веб-системы",
        "2.1 Общая архитектура",
        "2.2 Проектирование данных",
        "2.3 Проектирование API",
        "3 Реализация пользовательского интерфейса",
        "3.1 Главная страница и навигация",
        "3.2 Каталог и коллекции",
        "3.3 Карточка товара, избранное и корзина",
        "3.4 Блок выбора фактуры и цвета",
        "4 Реализация модуля визуализации интерьера",
        "4.1 Пользовательский сценарий",
        "4.2 Проверка изображения",
        "4.3 Сегментация стены",
        "4.4 Наложение текстуры",
        "4.5 Fallback и ручная коррекция",
        "5 Тестирование и оценка результата",
        "5.1 Функциональная проверка",
        "5.2 Адаптивность",
        "5.3 Ограничения",
        "6 Развертывание и сопровождение",
        "6.1 Запуск проекта",
        "6.2 Сопровождение",
        "ЗАКЛЮЧЕНИЕ",
        "СПИСОК ИСПОЛЬЗОВАННЫХ ИСТОЧНИКОВ",
        "ПРИЛОЖЕНИЕ А. Инструкция по запуску",
        "ПРИЛОЖЕНИЕ Б. Сценарий демонстрации",
    ]
    for item in items:
        p = doc.add_paragraph()
        p.paragraph_format.first_line_indent = Cm(0)
        p.paragraph_format.line_spacing = 1.2
        r = p.add_run(item)
        font(r, size=12)
    doc.add_page_break()


def intro(doc: Document) -> None:
    h1(doc, "ВВЕДЕНИЕ", page_break=False)
    for text in [
        "Интернет-магазин отделочных материалов не может ограничиваться карточками товара: покупателю нужно понять фактуру, цвет после покраски и вид покрытия в интерьере. Для жаккардовых обоев из кварцевой нити эта проблема особенно заметна, потому что материал продается под покраску и один и тот же рисунок меняется в зависимости от выбранного оттенка.",
        "Актуальность работы связана с разработкой веб-системы, которая объединяет каталог, коллекции, избранное, корзину, личный кабинет и модуль визуализации стены по фотографии комнаты. Такой сценарий снижает неопределенность выбора и позволяет показать покупателю не только товар, но и результат применения материала.",
        "Цель работы — спроектировать и реализовать веб-систему BauTex Design для выбора жаккардовых обоев, примерки цвета и автоматизированной визуализации покрытия на пользовательской фотографии интерьера.",
        "Для достижения цели решены следующие задачи: проанализированы пользовательские сценарии выбора обоев; спроектированы структура клиентского приложения и серверного API; реализованы каталог, карточка товара, избранное, корзина и личный кабинет; разработан отдельный сервис сегментации стен на базе SegFormer B0; реализованы проверка изображения, fallback-режим, наложение текстуры, сохранение результата и контейнеризация проекта.",
        "Объект исследования — процесс подбора отделочных материалов в веб-интерфейсе. Предмет исследования — методы проектирования веб-приложения и применения семантической сегментации для визуализации обоев на фотографии комнаты.",
        "Практическая значимость заключается в том, что полученная система может использоваться как демонстрационный прототип интернет-магазина: пользователь выбирает товар, проверяет его в интерьере, сохраняет результат и переходит к заказу. Теоретическая значимость состоит в сравнении подходов к выделению стен и обосновании выбора модели, пригодной для локального запуска.",
    ]:
        para(doc, text)


def chapter1(doc: Document) -> None:
    h1(doc, "1 Анализ предметной области и требований")
    h2(doc, "1.1 Характеристика предметной области")
    for text in [
        "BauTex Design представлен как сайт для продажи и подбора жаккардовых обоев из кварцевой нити. В отличие от обычного каталога, где товар описывается только фотографией и характеристиками, проект должен учитывать два измерения выбора: модель фактуры и цвет покраски.",
        "На практике покупатель проходит несколько шагов: знакомится с коллекциями, фильтрует товары, открывает карточку, выбирает цвет, добавляет позицию в избранное или корзину, а затем может загрузить фотографию комнаты и проверить выбранный материал на стене.",
    ]:
        para(doc, text)
    table(
        doc,
        "Таблица 1 – Основные пользовательские сценарии",
        ["Сценарий", "Реализация в проекте", "Результат для пользователя"],
        [
            ["Просмотр коллекций", "Отдельные страницы Basic, Loft, Geometry, Minimalism, Classic, Kids", "Пользователь быстро попадает в нужный стиль"],
            ["Поиск товара", "Фильтр по коллекции, стилю, цене, названию и артикулу", "Сокращается число неподходящих карточек"],
            ["Выбор цвета", "Палитра в карточке товара и в блоке примерки", "Покупатель видит, как меняется фактура"],
            ["Визуализация", "Загрузка фото, SegFormer/Fallback, режимы наложения", "Можно сравнить результат до заказа"],
            ["Сохранение", "Личный кабинет, избранное, корзина, проекты визуализации", "Сценарий не теряется после закрытия страницы"],
        ],
    )
    h2(doc, "1.2 Требования к системе")
    for text in [
        "Функциональные требования включают: отображение каталога и коллекций; поиск и фильтрацию товаров; открытие модальной карточки; добавление товара в избранное и корзину; регистрацию и вход; загрузку пользовательской фотографии; выбор режима сегментации и рендера; сохранение визуализации в аккаунт; просмотр информационных страниц и контактных данных.",
        "К нефункциональным требованиям отнесены адаптивность интерфейса, воспроизводимость локального запуска, отказоустойчивость при недоступности ML-сервиса, хранение секретов через переменные окружения, ограничение размера загружаемого файла и возможность контейнеризации.",
    ]:
        para(doc, text)
    h2(doc, "1.3 Выбор технологического стека")
    table(
        doc,
        "Таблица 2 – Использованные технологии",
        ["Уровень", "Технологии", "Назначение"],
        [
            ["Frontend", "React 19, TypeScript, React Router, Redux Toolkit", "Маршрутизация, состояние пользователя и избранного, компонентный интерфейс"],
            ["Backend", "Node.js, Express, SQLite, JWT, bcryptjs, multer, sharp", "API, авторизация, загрузка изображений, обработка файлов, хранение данных"],
            ["ML-сервис", "FastAPI, PyTorch, Hugging Face Transformers, SegFormer B0", "Семантическая сегментация стен"],
            ["Инфраструктура", "Webpack, Docker Compose", "Сборка, dev-сервер, воспроизводимый запуск нескольких сервисов"],
        ],
    )


def chapter2(doc: Document) -> None:
    h1(doc, "2 Проектирование веб-системы")
    h2(doc, "2.1 Общая архитектура")
    para(doc, "Система разделена на три части: клиентское React-приложение, Express-сервер и отдельный FastAPI-сервис сегментации. Такое разделение позволяет развивать интерфейс и ML-модуль независимо: при недоступности нейросети веб-приложение сохраняет базовые функции каталога и может использовать fallback-режим.")
    table(
        doc,
        "Таблица 3 – Архитектурные компоненты",
        ["Компонент", "Файлы проекта", "Ответственность"],
        [
            ["Клиент", "src/App.tsx, src/components, src/pages", "Маршруты, экраны, состояние, взаимодействие пользователя"],
            ["API", "server/index.js, server/routes/*.js", "Авторизация, каталог, корзина, избранное, визуализация"],
            ["База данных", "server/database.js, server/migrations/001_core_schema.js", "Пользователи, заказы, товары, проекты, метрики"],
            ["Сегментация", "ml_service/app.py", "Загрузка модели и возврат PNG-маски стены"],
        ],
    )
    code_excerpt(
        doc,
        "Листинг 1 – Фрагмент маршрутизации клиентского приложения",
        excerpt_between(read("src/App.tsx"), "<Routes>", "</Routes>"),
        34,
    )
    h2(doc, "2.2 Проектирование данных")
    para(doc, "SQLite выбран как локальное хранилище для прототипа: он не требует отдельного сервера БД, легко переносится в Docker volume и достаточен для демонстрационного интернет-магазина. Схема вынесена в миграцию, поэтому при запуске сервер сам создает недостающие таблицы и добавляет новые поля без ручной подготовки базы.")
    code_excerpt(
        doc,
        "Листинг 2 – Фрагмент миграции SQLite",
        excerpt_between(read("server/migrations/001_core_schema.js"), "CREATE TABLE IF NOT EXISTS users", "CREATE TABLE IF NOT EXISTS visualization_projects"),
        38,
    )
    h2(doc, "2.3 Проектирование API")
    para(doc, "API строится вокруг ресурсов: auth, catalog, cart, favorites, visualization, mlMetrics и aiChat. Для пользовательских действий, связанных с аккаунтом, используется JWT. Для изображений применяется multer, затем sharp получает метаданные и выполняет подготовку файлов к обработке.")
    table(
        doc,
        "Таблица 4 – Основные API-операции",
        ["Маршрут", "Метод", "Назначение"],
        [
            ["/api/auth/login", "POST", "Вход пользователя и выдача JWT"],
            ["/api/catalog/products", "GET", "Получение администрируемых товаров"],
            ["/api/favorites", "GET/POST/DELETE", "Синхронизация избранного"],
            ["/api/cart", "GET/POST/PATCH/DELETE", "Работа с корзиной авторизованного пользователя"],
            ["/api/visualize", "POST", "Создание визуализации по фото, маске, фактуре и цвету"],
            ["/api/visualize/check", "POST", "Проверка размера и качества загруженного изображения"],
        ],
    )


def chapter3(doc: Document) -> None:
    h1(doc, "3 Реализация пользовательского интерфейса")
    h2(doc, "3.1 Главная страница и навигация")
    para(doc, "Главная страница реализована как последовательность самостоятельных блоков: Hero, Features, Slider, FabricsSection, LocationMap и Reviews. Это позволяет менять визуальные блоки без изменения маршрутизации и без смешивания логики каталога с презентационной частью.")
    image(doc, SCREENS / "home.png", "Рисунок 1 – Главная страница BauTex Design", 15.5)
    h2(doc, "3.2 Каталог и коллекции")
    para(doc, "Каталог объединяет товары из локального набора данных и позиции, полученные с backend. Если администратор добавляет товар с тем же идентификатором, серверная версия перекрывает статическую. Такой подход сохраняет демо-данные и одновременно дает возможность расширять ассортимент.")
    image(doc, SCREENS / "catalog-desktop.png", "Рисунок 2 – Каталог с фильтрами и карточками товаров", 15.5)
    code_excerpt(
        doc,
        "Листинг 3 – Фильтрация и сортировка товаров в каталоге",
        excerpt_between(read("src/pages/Catalog/Catalog.tsx"), "const filteredItems = useMemo", "const handleToggleFavorite"),
        34,
    )
    h2(doc, "3.3 Карточка товара, избранное и корзина")
    para(doc, "Карточка товара открывается поверх каталога. В ней пользователь выбирает цвет покраски, количество рулонов, добавляет товар в корзину или избранное. Для корзины реализована двойная стратегия: локальное сохранение в localStorage работает даже без входа, а после авторизации данные дополнительно синхронизируются с сервером.")
    code_excerpt(
        doc,
        "Листинг 4 – Добавление товара в локальную и серверную корзину",
        excerpt_between(read("src/components/ProductModal.tsx"), "const addToCart = async", "return ("),
        42,
    )
    h2(doc, "3.4 Блок выбора фактуры и цвета")
    para(doc, "Блок FabricsSection связывает выбор коллекции, конкретной фактуры и оттенка. При смене цвета меняется фон секции и слой поверх текстуры. Такой экран используется не как декоративная галерея, а как быстрый способ показать зависимость материала от покраски.")
    code_excerpt(
        doc,
        "Листинг 5 – Расчет фонового оттенка блока примерки",
        excerpt_between(read("src/components/FabricsSection.tsx"), "const getTintedBackground", "const FabricsSection"),
        30,
    )


def chapter4(doc: Document) -> None:
    h1(doc, "4 Реализация модуля визуализации интерьера")
    h2(doc, "4.1 Пользовательский сценарий")
    para(doc, "Страница визуализации принимает фотографию комнаты, выбранную фактуру, цвет покраски, масштаб принта, интенсивность цвета и режим наложения. Пользователь может выбрать SegFormer B0 или fallback, увидеть предупреждения о качестве фото и при необходимости перейти к ручной коррекции маски.")
    image(doc, SCREENS / "visualization-desktop.png", "Рисунок 3 – Интерфейс модуля визуализации интерьера", 15.5)
    h2(doc, "4.2 Проверка изображения")
    para(doc, "Перед обработкой сервер анализирует файл через sharp: проверяет размеры, соотношение сторон и вес. Это не запрещает обработку, но возвращает предупреждения, чтобы интерфейс мог подсказать пользователю, почему результат может быть хуже.")
    code_excerpt(
        doc,
        "Листинг 6 – Проверка пользовательской фотографии",
        excerpt_between(read("server/routes/visualization.js"), "router.post('/api/visualize/check'", "async function createFallbackWallpaper"),
        38,
    )
    h2(doc, "4.3 Сегментация стены")
    para(doc, "Для автоматического выделения стены используется модель nvidia/segformer-b0-finetuned-ade-512-512. FastAPI-сервис загружает модель один раз, определяет идентификаторы классов wall, выполняет инференс и возвращает бинарную PNG-маску. Постобработка сглаживает маску медианным фильтром и Gaussian Blur, затем порог отделяет стену от остальных областей.")
    code_excerpt(
        doc,
        "Листинг 7 – Endpoint FastAPI для сегментации стены",
        excerpt_between(read("ml_service/app.py"), "@app.post(\"/segment/wall\")", None),
        40,
    )
    h2(doc, "4.4 Наложение текстуры")
    para(doc, "Backend получает маску и накладывает на отмеченную область подготовленную текстуру обоев. Для выбранного цвета сервер формирует tinted wallpaper buffer, а затем режимы catalog, realistic и contrast управляют тем, насколько сохраняются тени исходной фотографии и насколько выражен рисунок материала.")
    code_excerpt(
        doc,
        "Листинг 8 – Подготовка цветной фактуры и наложение по маске",
        excerpt_between(read("server/routes/visualization.js"), "async function createTintedWallpaperBuffer", "async function createHeuristicWallMask"),
        44,
    )
    h2(doc, "4.5 Fallback и ручная коррекция")
    para(doc, "Нейросетевой режим не является единственным способом завершить сценарий. Если ML-сервис недоступен или маска покрывает слишком малую/слишком большую часть изображения, система может перейти к эвристической маске и предложить пользователю ручное уточнение кистью. Это важно для демонстрационного прототипа: сбой модели не должен блокировать весь пользовательский путь.")


def chapter5(doc: Document) -> None:
    h1(doc, "5 Тестирование и оценка результата")
    h2(doc, "5.1 Функциональная проверка")
    para(doc, "Функциональная проверка выполнялась по пользовательским сценариям: открытие главной страницы, переходы по меню, фильтрация каталога, открытие карточки, добавление товара в избранное и корзину, вход в аккаунт, загрузка фотографии, выбор фактуры и запуск визуализации.")
    table(
        doc,
        "Таблица 5 – Проверенные сценарии",
        ["Сценарий", "Ожидаемое поведение", "Фактический результат"],
        [
            ["Каталог", "Фильтры сужают список товаров", "Товары пересчитываются через useMemo без перезагрузки страницы"],
            ["Избранное", "Состояние видно после клика", "Redux Toolkit хранит список id, сервер синхронизирует его по JWT"],
            ["Корзина", "Товар сохраняется даже без авторизации", "localStorage используется как первый уровень, backend — как синхронизация"],
            ["Фото комнаты", "Некорректный файл не проходит обработку", "multer ограничивает тип и размер, sharp возвращает предупреждения"],
            ["Сегментация", "Сервис возвращает маску стены", "FastAPI возвращает PNG-маску, backend анализирует покрытие"],
        ],
    )
    h2(doc, "5.2 Адаптивность")
    para(doc, "Интерфейс проверялся в desktop- и mobile-ширинах. На малой ширине навигация сворачивается, hero-блок и карточки перестраиваются, а примерочная сохраняет последовательность выбора: сначала модель, затем цвет и действие. Это позволяет использовать сайт с телефона без горизонтальной прокрутки.")
    image(doc, SCREENS / "mobile-home.png", "Рисунок 4 – Проверка главной страницы в мобильном viewport", 7.0)
    h2(doc, "5.3 Ограничения")
    for text in [
        "Качество сегментации зависит от исходной фотографии. Сложные интерьеры с зеркалами, плотной мебелью у стены, сильными тенями или нестандартным ракурсом могут давать неполную маску. Поэтому в проекте сохранены fallback-режим и ручная кисть.",
        "В прототипе SQLite подходит для локальной демонстрации и контейнерного запуска, но при промышленной эксплуатации потребуется перейти на отдельную СУБД, добавить миграции с версионированием, очереди для тяжелой обработки изображений и хранение файлов в объектном хранилище.",
    ]:
        para(doc, text)


def chapter6(doc: Document) -> None:
    h1(doc, "6 Развертывание и сопровождение")
    h2(doc, "6.1 Запуск проекта")
    para(doc, "Проект можно запускать как по отдельным сервисам, так и через Docker Compose. В compose-файле описаны frontend, backend и ml. Backend зависит от ML-сервиса, а frontend зависит от backend. Для данных загрузок и результатов используются volume, поэтому контейнеры можно пересоздавать без потери пользовательских файлов.")
    code_excerpt(
        doc,
        "Листинг 9 – Состав сервисов Docker Compose",
        read("docker-compose.yml"),
        44,
    )
    h2(doc, "6.2 Сопровождение")
    para(doc, "Структура проекта поддерживает расширение: новые коллекции добавляются через CATALOG_ITEMS или API продуктов; новые информационные страницы подключаются как lazy routes; ML-модель может быть заменена через переменную SEGFORMER_MODEL; параметры backend задаются через .env.")
    table(
        doc,
        "Таблица 6 – Направления дальнейшего развития",
        ["Направление", "Что требуется изменить"],
        [
            ["Админ-панель каталога", "Добавить формы редактирования products и collections поверх существующей схемы"],
            ["Очередь ML-задач", "Вынести обработку изображений в фоновые задания и показывать статус"],
            ["Метрики качества", "Расширить test_data/segmentation и сохранять результаты прогонов в ml_metric_runs"],
            ["Коммерческая эксплуатация", "Заменить SQLite на PostgreSQL и вынести файлы в S3-совместимое хранилище"],
        ],
    )


def conclusion(doc: Document) -> None:
    h1(doc, "ЗАКЛЮЧЕНИЕ")
    for text in [
        "В результате работы разработана веб-система BauTex Design, объединяющая каталог жаккардовых обоев, карточку товара, избранное, корзину, личный кабинет, информационные страницы и модуль визуализации интерьера.",
        "Практическая часть включает клиентское приложение на React и TypeScript, backend на Express с SQLite-хранилищем и отдельный FastAPI-сервис сегментации стен. Для визуализации реализованы автоматический режим SegFormer B0, fallback, проверка входного фото, наложение фактуры с выбранным цветом и сохранение результата.",
        "Поставленные задачи выполнены: определены требования и пользовательские сценарии; спроектирована архитектура; реализованы основные функции сайта; добавлены механизмы обработки изображений и контейнеризация; проведена функциональная и адаптивная проверка. Полученный прототип можно использовать как основу для дальнейшего развития интернет-магазина с визуальной примеркой обоев.",
    ]:
        para(doc, text)


def sources(doc: Document) -> None:
    h1(doc, "СПИСОК ИСПОЛЬЗОВАННЫХ ИСТОЧНИКОВ")
    items = [
        "ГОСТ 7.32-2017. СИБИД. Отчет о научно-исследовательской работе. Структура и правила оформления.",
        "ГОСТ Р 7.0.100-2018. СИБИД. Библиографическая запись. Библиографическое описание.",
        "React Documentation. URL: https://react.dev/ (дата обращения: 10.05.2026).",
        "Redux Toolkit Documentation. URL: https://redux-toolkit.js.org/ (дата обращения: 10.05.2026).",
        "React Router Documentation. URL: https://reactrouter.com/ (дата обращения: 10.05.2026).",
        "Express.js Documentation. URL: https://expressjs.com/ (дата обращения: 10.05.2026).",
        "SQLite Documentation. URL: https://www.sqlite.org/docs.html (дата обращения: 10.05.2026).",
        "Sharp Documentation. URL: https://sharp.pixelplumbing.com/ (дата обращения: 10.05.2026).",
        "FastAPI Documentation. URL: https://fastapi.tiangolo.com/ (дата обращения: 10.05.2026).",
        "PyTorch Documentation. URL: https://pytorch.org/docs/stable/index.html (дата обращения: 10.05.2026).",
        "Hugging Face Transformers Documentation. URL: https://huggingface.co/docs/transformers/ (дата обращения: 10.05.2026).",
        "Xie E., Wang W., Yu Z., Anandkumar A., Alvarez J. M., Luo P. SegFormer: Simple and Efficient Design for Semantic Segmentation with Transformers. URL: https://arxiv.org/abs/2105.15203 (дата обращения: 10.05.2026).",
        "Docker Compose Documentation. URL: https://docs.docker.com/compose/ (дата обращения: 10.05.2026).",
    ]
    for i, item in enumerate(items, 1):
        para(doc, f"{i}. {item}", indent=False)


def appendices(doc: Document) -> None:
    h1(doc, "ПРИЛОЖЕНИЕ А. Инструкция по запуску")
    for text in [
        "1. Создать файл server/.env по образцу server/.env.example и задать JWT_SECRET, ADMIN_EMAIL, ADMIN_PASSWORD, FRONTEND_ORIGIN и SEGFORMER_API_URL.",
        "2. Установить зависимости командой npm install в корне проекта и npm install в каталоге server.",
        "3. Запустить backend: npm run server:start. Проверить http://localhost:3003/health.",
        "4. Запустить frontend: npm start. Открыть http://localhost:3001.",
        "5. Для ML-сервиса установить зависимости из ml_service/requirements.txt и выполнить npm run ml:start либо использовать docker compose up --build.",
    ]:
        para(doc, text, indent=False)
    h1(doc, "ПРИЛОЖЕНИЕ Б. Сценарий демонстрации")
    for text in [
        "1. Открыть главную страницу, показать hero-блок, меню и блок подбора фактуры/цвета.",
        "2. Перейти в каталог, применить фильтр по коллекции, открыть карточку товара.",
        "3. Выбрать цвет, добавить товар в избранное и корзину, показать изменение состояния кнопок.",
        "4. Открыть страницу визуализации, загрузить фото комнаты, выбрать фактуру и режим SegFormer.",
        "5. Показать результат, маску стены, предупреждения о качестве фото и fallback-режим.",
    ]:
        para(doc, text, indent=False)


def main() -> None:
    doc = Document()
    configure(doc)
    cover(doc)
    toc(doc)
    intro(doc)
    chapter1(doc)
    chapter2(doc)
    chapter3(doc)
    chapter4(doc)
    chapter5(doc)
    chapter6(doc)
    conclusion(doc)
    sources(doc)
    appendices(doc)
    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    main()
